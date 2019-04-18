import json

import pymysql
from docx import Document
import os
import urllib.request
import docx
from docx.enum.text import WD_BREAK

from docx.shared import Inches

#黄晨阳
#郭文俊
#刘金龙
#陈鹏程

#服务器端的ip地址
# serverIP = "10.101.214.26"
serverIP = "10.56.33.12"

#在获取服务器端的图片时，使用的地址前缀
serverPrePath = "http://{serverIP}:8080".format(serverIP=serverIP)

#在读取数据库中的图片时，使用的地址
mysqlPrePath = serverIP

#生成的常见错题，不带答案的版本
alwaysMistakeQuestion = r"amq.docx"

#获得图片后面的题干内容
def getTiGanSuffix(tigan) :
    # print(tigan[0:tigan.index("<")])
    return tigan[tigan.index(">")+1:]

#获得图片前面的题干内容
def getTiGanPrefix(tigan) :
    # print(tigan[0:tigan.index("<")])
    return tigan[0:tigan.index("<")]



# 从数据库中读取指定的题（包含选项）
# 并把题干和选项写入到指定的word文件中
def readMySQLAndWriteQuestToWord(document, qsnId) :

    aChoiceFile = open(r"./A.png", "wb+")
    bChoiceFile = open(r"./B.png", "wb+")
    cChoiceFile = open(r"./C.png", "wb+")
    dChoiceFile = open(r"./D.png", "wb+")
    titleFile = open(r"title.png", "wb+")

    choiceFileList = [ aChoiceFile, bChoiceFile, cChoiceFile, dChoiceFile ]

    mcursor = db.cursor()

    # print("qsnId = " + qsnId)

    mcursor.execute("SELECT * FROM mta_exam_shiti_danxuan_data where qsnId = {qsnId}".format(qsnId=qsnId))

    mydatas = mcursor.fetchall()

    for dat in mydatas:
        #读取题干
        print(dat[1])
        #判断题干中是否有图片
        # if分支表示有图片，则将图片写入到文件
        if dat[1].find("img src") != -1 :
            photoPath = getPhotoPath (dat[1])
            # 获取题干
            tiGanPre = getTiGanPrefix(dat[1])
            # print(tiGanPre)

            tiGanSuf = getTiGanSuffix(dat[1])
            # print(tiGanSuf)

            pic = readPhoto(serverPrePath, photoPath)

            writePhoto(titleFile, pic)

            mPara = document.add_paragraph(tiGanPre, style="List Number")
            run = mPara.add_run()

            # writeToWordPhoto(document, titleFile)
            writeToWordPhotoRun(run, titleFile)

            mPara.add_run(tiGanSuf)

            document.save(alwaysMistakeQuestion)

        else :
            # 获取题干


            #如果没有图片，则直接将题干写入到文件中
            # para = writeToWordParagraph(document, dat[1], alwaysMistakeQuestion)
            mPara = document.add_paragraph(dat[1], style="List Number")
            document.save(alwaysMistakeQuestion)


        #从mysql服务器中读取指定题的选项
        print(dat)
        # print(dat[2])

        choicePreList = ["A:","B:","C:","D:"]

        for i in range (1,5) :

            mpara = document.add_paragraph(choicePreList[i-1])

            if dat[2].find("img src") != -1 :
                run = mpara.add_run()

                print("选项中有图片")

                dictChoices = json.loads(dat[2])

                choicePath = getPhotoPath(dictChoices["{pos}".format(pos=i)]["xx"])


                try :
                    print("choicePath = " + choicePath)
                    writePhoto(choiceFileList[i-1], readPhoto(serverPrePath, choicePath))
                    # writeToWordPhoto(document, choiceFileList[i-1])
                    writeToWordPhotoRunWithWidth(run, choiceFileList[i-1], 1)
                except TypeError as e:
                    print("有异常发生")
                    # writeToWordParagraph(document, dictChoices["{pos}".format(pos = i)]["xx"], alwaysMistakeQuestion)
                    writeToWordParagraphRun(run, dictChoices["{pos}".format(pos = i)]["xx"], alwaysMistakeQuestion)


            else :

                print("选项中没有图片")

                # 如果没有图片，则直接将选项写入到文件中
                dictChoices = json.loads(dat[2])

                run = mpara.add_run(dictChoices["{pos}".format(pos = i)]["xx"])

                # writeToWordParagraph(document, dictChoices["{pos}".format(pos = i)]["xx"], alwaysMistakeQuestion)
                document.save(alwaysMistakeQuestion)

        # print(dat[3])

        dpara = document.add_paragraph("答案：")
        dpara.add_run(choicePreList[dat[3]-1])

        #关闭所有已经打开的文件
        # 防止因为文件忘记关闭导致的内存泄漏
        aChoiceFile.close()
        bChoiceFile.close()
        cChoiceFile.close()
        dChoiceFile.close()

        document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        document.save(alwaysMistakeQuestion)


# 将从数据库中获取的字符串通过切片的方式，拿到里面的图片路径
# 该路径是在服务器端保存图片的绝对路径
# 需要使用地址前缀加绝对路径的方式来得到真正可以使用的网络路径
# 该网络路径最终可用于获取指定的图片

class MyNotPhotoException (ValueError) :
    def showMsg(self) :
        return "is a normal string"

def getPhotoPath(dat) :

    try :
        photoPath = dat[dat.index("'")+1:dat.rindex("'")]
    except ValueError:
        return -1

    return photoPath

#向word文件中写入文本
#没有关闭document流对象
def writeToWordParagraphRun(document, content, targetFile) :
    document.paragraphs[0].add_run("content")

#向word文件中写入文本
#没有关闭document流对象
def writeToWordParagraph(document, content, targetFile) :
    para = document.add_paragraph(content)
    document.save(targetFile)

    return para


#向word文件中写入图片
#没有关闭document流对象
def writeToWordPhotoRunWithWidth(run, titleFile, minches) :
    # document.paragraphs
    run.add_picture(titleFile, width=Inches(minches))

#向word文件中写入图片
#没有关闭document流对象
def writeToWordPhotoRun(run, titleFile) :
    # document.paragraphs
    run.add_picture(titleFile, width=Inches(2))

#向word文件中写入图片
#没有关闭document流对象
def writeToWordPhoto(document, titleFile) :
    # document.paragraphs
    document.add_picture(titleFile)
    document.save(alwaysMistakeQuestion)

#将从服务器端获取的内容写入文件
#此处这个函数用于将图片写入到指定文件
#然后再讲该图片添加到word文档中。
#注意此处没有执行流的关闭操作
#   要防止因为流操作造成的内存泄漏
def writePhoto(fileobj, writeContent):
    fileobj.write(writeContent)

#返回读取的图片
def readPhoto(prePath, realPath) :
    try :
        response = urllib.request.urlopen(prePath+realPath)
        pic = response.read()
    except TypeError:
        return -1

    return pic

#读取错题序号文件
#从该文件中读取时，因为是复制的网页中的内容，
# 因此是表格形式
def readWrongWordFile(filepath) :
    wrongQuestList = list()

    documentWrong = docx.Document(filepath)

    tbls = documentWrong.tables
    tbl = tbls[0]

    for i in range(0, len(tbl.rows)):
        wrongQuestList.append( tbl.cell(i,1).text )

    return wrongQuestList

if __name__ == '__main__':

    wrongQuestList = readWrongWordFile(r"./wrongquest.docx")
    print(wrongQuestList)

    # print(wrongQuestList)

    if os.path.exists(alwaysMistakeQuestion):
        document = Document(alwaysMistakeQuestion)
    else :
        document = Document()

    db = pymysql.connect(
        mysqlPrePath, #数据库对应的地址
        "root", # 数据库登录用户名
        "root", # 数据库登录密码
        "2015mtaks3" # 数据库名称
    )

    for wr in wrongQuestList:
        readMySQLAndWriteQuestToWord(document, wr)
